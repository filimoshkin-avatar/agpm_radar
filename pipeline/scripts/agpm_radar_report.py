#!/usr/bin/env python3
"""Generate a weekly AgPM radar report in Markdown and DOCX."""

from __future__ import annotations

import argparse
import email.utils
import hashlib
import html
import json
import re
import sys
import urllib.parse
from collections import Counter, defaultdict
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_WIKI = ROOT / "knowledge/agpm-radar"
PERIMETERS = [
    ("far", "Дальний периметр", "бизнес-агенты, агентные рабочие среды, платформы, policy/market-сигналы и институциональные события, которые могут изменить управленческий контур бизнеса"),
    ("middle", "Средний периметр", "агентное управление в бизнесе: агентные бизнес-процессы, операционные модели, исследования, внедрения и кейсы"),
    ("near", "Близкий периметр", "агентное управление проектами, портфелями и проектными офисами"),
]
SECTION_ANCHORS = {
    "Что важно для AgPM": "section_agpm_implications",
    "Разбор AI Agents Directory": "section_aiagents_directory",
    "Дальний периметр": "section_far",
    "Средний периметр": "section_middle",
    "Близкий периметр": "section_near",
}
TOC_LINKS = [
    ("Что важно для AgPM", "section_agpm_implications"),
    ("Разбор AI Agents Directory", "section_aiagents_directory"),
    ("Дальний периметр", "section_far"),
    ("Средний периметр", "section_middle"),
    ("Близкий периметр", "section_near"),
]
REPORT_LINK_RE = re.compile(r"^\s*(?:-\s*)?Ссылка:\s+(\S+)", re.MULTILINE)
REPORT_DATE_RE = re.compile(r"AgPM_(?:daily|weekly)_radar_(\d{4}-\d{2}-\d{2})\.md$")
FULLTEXT_CACHE_DIRNAME = "source-fulltext"
DAILY_REPORT_LIMIT = 10
DAILY_DEFERRED_QUEUE = "daily-deferred.jsonl"
TRACKING_PARAMS = {
    "utm_source",
    "utm_medium",
    "utm_campaign",
    "utm_content",
    "utm_term",
    "fbclid",
    "gclid",
    "yclid",
    "mc_cid",
    "mc_eid",
}
LOCALE_SEGMENTS = {
    "en",
    "es",
    "nl",
    "de",
    "fr",
    "it",
    "pt",
    "pt-br",
    "zh",
    "zh-cn",
    "zh-tw",
    "cn",
    "ja",
    "ko",
    "ru",
    "pl",
    "tr",
}
LANGUAGE_BY_LOCALE = {
    "en": "английском",
    "es": "испанском",
    "nl": "нидерландском",
    "de": "немецком",
    "fr": "французском",
    "it": "итальянском",
    "pt": "португальском",
    "pt-br": "португальском",
    "zh": "китайском",
    "zh-cn": "китайском",
    "zh-tw": "китайском",
    "cn": "китайском",
    "ja": "японском",
    "ko": "корейском",
    "ru": "русском",
    "pl": "польском",
    "tr": "турецком",
}

AGENT_TERMS = [
    "ai agent",
    "ai agents",
    "agentic",
    "autonomous agent",
    "business agent",
    "agent ai",
    "agent economy",
    "intelligent agent",
    "intelligent agents",
    "ai-агент",
    "ai агентов",
    "ai-агентов",
    "multi-agent",
    "agent orchestration",
    "ai agent framework",
    "agentforce",
    "agentic security",
    "computer use",
    "perplexity computer",
    "openclaw",
    "агентн",
    "агентов",
    "агента",
    "агенты",
    "ии-агент",
    "мультиагент",
]

MANAGEMENT_TERMS = [
    "management",
    "governance",
    "orchestration",
    "workflow",
    "business process",
    "operations",
    "operating model",
    "enterprise",
    "organization",
    "accountability",
    "policy",
    "policy framework",
    "government",
    "government framework",
    "regulation",
    "regulatory",
    "public sector",
    "state policy",
    "industrial policy",
    "agent economy",
    "joint venture",
    "ai services",
    "permissions",
    "security",
    "observability",
    "durable execution",
    "execution histories",
    "recovery",
    "risk",
    "compliance",
    "human in the loop",
    "human-in-the-loop",
    "pmo",
    "project management",
    "portfolio management",
    "project portfolio",
    "customer support",
    "sales",
    "crm",
    "управлен",
    "оркестр",
    "бизнес-процесс",
    "бизнес процесс",
    "операцион",
    "корпоратив",
    "инфраструктур",
    "ии-сервис",
    "ии сервис",
    "ai-сервис",
    "ai сервис",
    "совместное предприятие",
    "технологическ",
    "платформ",
    "инвестиц",
    "ответствен",
    "риск",
    "контроль",
    "портфел",
    "проектный офис",
]

NEAR_TERMS = [
    "ai pmo",
    "pmo",
    "project management",
    "portfolio management",
    "project portfolio",
    "проектный офис",
    "управлени проект",
    "управление проект",
    "портфел",
]

MIDDLE_TERMS = [
    "business process",
    "workflow",
    "operations",
    "operating model",
    "enterprise",
    "customer support",
    "sales",
    "crm",
    "agent orchestration",
    "business agent",
    "policy framework",
    "government framework",
    "state policy",
    "public sector",
    "agent economy",
    "intelligent agents",
    "joint venture",
    "ai services",
    "бизнес-процесс",
    "бизнес процесс",
    "операцион",
    "корпоратив",
    "инфраструктур",
    "ии-сервис",
    "ии сервис",
    "ai-сервис",
    "ai сервис",
    "совместное предприятие",
    "технологическ",
    "платформ",
    "инвестиц",
]

FAR_TERMS = [
    "agentforce",
    "business agent",
    "perplexity computer",
    "openclaw",
    "computer use",
    "workspace studio",
    "agent workspace",
    "ai agent orchestration",
    "autonomous ai platform",
]

EXCLUDE_TERMS = [
    "code",
    "claude code",
    "coding",
    "developer",
    "github",
    "repository",
    "repo",
    "programming",
    "python",
    "javascript",
    "sdk",
    "cli",
    "terminal",
    "tutorial",
    "workshop",
    "consulting",
    "how to build",
    "build a",
    "firecrawl",
    "agno",
    "gpt-4o",
    "devops",
    "software engineering",
    "pytest",
    "bug",
    "bugs",
    "subagents",
    "баг",
    "код",
    "разработ",
    "программ",
    "devops",
    "lsp",
    "rag",
    "ocr",
    "video",
    "filmmaking",
    "creator",
    "deepfake",
    "k-pop",
    "game",
    "mmo",
    "education",
    "bible",
    "powerpoint maker",
    "hardware",
    "gpu",
    "avx",
    "кино",
    "видео",
    "программирован",
    "разработчик",
    "кодинг",
    "репозитор",
    "agency",
    "агентств",
    "trading",
    "credit card",
    "wallet",
    "defi",
    "onchain",
    "трейдинг",
    "кредитн",
    "карта",
    "кошелек",
    "кошелёк",
    "ончен",
    "покупк",
]

ACTION_TERMS = [
    "execute",
    "executes",
    "take actions",
    "taking actions",
    "automate",
    "automation",
    "autonomous workflow",
    "autonomous workflows",
    "update timelines",
    "assign tasks",
    "flag resource conflicts",
    "schedule risks",
    "reporting",
    "tracking",
    "risk detection",
    "resource",
    "tasks",
    "appointments",
    "closing sales",
    "customer service",
    "customer support",
    "выполня",
    "автоматизац",
    "координац",
    "отчет",
    "отчёт",
    "контроль",
    "поручен",
    "рисков",
    "ресурс",
]

WEAK_SIGNAL_TERMS = [
    "swarm",
    "bees",
    "ants",
    "fish",
    "cybersecurity",
    "smart cities",
    "логистические сети",
    "умные города",
    "пчёлы",
    "муравьи",
    "рыбы",
    "кибербезопас",
]

STRONG_BUSINESS_PRODUCTS = [
    "meta business agent",
    "whatsapp business",
    "agentforce",
    "workspace studio",
    "perplexity computer",
    "ai agent orchestration",
    "monday.com welcomes ai agents",
    "respond.io",
]

STRONG_MANAGEMENT_CONTEXT = [
    "governance",
    "business process",
    "business operations",
    "operating model",
    "enterprise operations",
    "customer support",
    "sales",
    "crm",
    "policy",
    "policy framework",
    "government",
    "government framework",
    "regulation",
    "regulatory",
    "public sector",
    "state policy",
    "industrial policy",
    "agent economy",
    "joint venture",
    "ai services",
    "permissions",
    "security",
    "observability",
    "durable execution",
    "execution histories",
    "recovery",
    "compliance",
    "risk management",
    "project management",
    "portfolio management",
    "project portfolio",
    "pmo",
    "agentic pmo",
    "ai pmo",
    "управление проект",
    "проектный офис",
    "портфель",
    "бизнес-процесс",
    "операцион",
    "корпоратив",
    "инфраструктур",
    "ии-сервис",
    "ии сервис",
    "ai-сервис",
    "ai сервис",
    "совместное предприятие",
    "технологическ",
    "платформ",
    "инвестиц",
    "ответствен",
]

FINANCIAL_AGENT_TERMS = [
    "trading",
    "credit card",
    "wallet",
    "defi",
    "onchain",
    "трейдинг",
    "кредитн",
    "кошелек",
    "кошелёк",
    "покупк",
]

PROMOTIONAL_URL_HINTS = [
    "/product/",
    "/products/",
    "/solutions/",
    "/services/",
    "/platform-enterprise",
    "/agent-platform-enterprise",
]

PROMOTIONAL_TEXT_TERMS = [
    "request demo",
    "book a demo",
    "contact sales",
    "schedule demo",
    "запросить демо",
    "заказать демо",
    "связаться с продажами",
]


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def parse_dt(value: str | None) -> datetime | None:
    if not value:
        return None
    value = value.strip()
    partial = re.fullmatch(r"(\d{4})-(\d{2}|\?\?)-(\d{2}|\?\?)", value)
    if partial:
        year = int(partial.group(1))
        month = 1 if partial.group(2) == "??" else int(partial.group(2))
        day = 1 if partial.group(3) == "??" else int(partial.group(3))
        return datetime(year, month, day, tzinfo=timezone.utc)
    relative = re.fullmatch(r"(\d+)\s+(day|days|week|weeks|month|months|year|years)\s+ago", value.lower())
    if relative:
        amount = int(relative.group(1))
        unit = relative.group(2)
        days = amount
        if unit.startswith("week"):
            days = amount * 7
        elif unit.startswith("month"):
            days = amount * 30
        elif unit.startswith("year"):
            days = amount * 365
        return utc_now() - timedelta(days=days)
    try:
        return email.utils.parsedate_to_datetime(value).astimezone(timezone.utc)
    except Exception:
        pass
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).astimezone(timezone.utc)
    except Exception:
        pass
    for fmt in ("%A, %B %d, %Y", "%a, %b %d, %Y", "%B %d, %Y", "%b %d, %Y"):
        try:
            return datetime.strptime(value, fmt).replace(tzinfo=timezone.utc)
        except Exception:
            pass
    return None


def load_materials(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    items: list[dict[str, Any]] = []
    with path.open("r", encoding="utf-8") as stream:
        for line in stream:
            if line.strip():
                items.append(json.loads(line))
    return items


def canonicalize_report_url(url: str | None) -> str:
    if not url:
        return ""
    url = url.strip().rstrip(".,);")
    parsed = urllib.parse.urlsplit(url)
    scheme = parsed.scheme.lower() or "https"
    netloc = parsed.netloc.lower()
    if netloc.startswith("www."):
        netloc = netloc[4:]
    query = urllib.parse.parse_qsl(parsed.query, keep_blank_values=False)
    query = [(key, value) for key, value in query if key.lower() not in TRACKING_PARAMS]
    path = re.sub(r"/{2,}", "/", parsed.path or "/")
    if path != "/":
        path = path.rstrip("/")
    return urllib.parse.urlunsplit((scheme, netloc, path, urllib.parse.urlencode(query, doseq=True), ""))


def strip_locale_path(path: str) -> str:
    parts = [part for part in path.split("/") if part]
    if not parts:
        return path or "/"
    first = parts[0].lower()
    if first in LOCALE_SEGMENTS or re.fullmatch(r"[a-z]{2}-[a-z]{2}", first):
        parts = parts[1:]
    return "/" + "/".join(parts) if parts else "/"


def locale_neutral_report_url(url: str | None) -> str:
    canonical = canonicalize_report_url(url)
    if not canonical:
        return ""
    parsed = urllib.parse.urlsplit(canonical)
    path = strip_locale_path(parsed.path)
    if path != "/":
        path = path.rstrip("/")
    return urllib.parse.urlunsplit((parsed.scheme, parsed.netloc, path, parsed.query, ""))


def language_from_url(url: str | None) -> str | None:
    if not url:
        return None
    parsed = urllib.parse.urlsplit(url)
    parts = [part for part in parsed.path.split("/") if part]
    if not parts:
        return None
    locale = parts[0].lower()
    if re.fullmatch(r"[a-z]{2}-[a-z]{2}", locale):
        locale = locale.split("-")[0]
    return LANGUAGE_BY_LOCALE.get(locale)


def detected_language(item: dict[str, Any]) -> str | None:
    language = language_from_url(item.get("url")) or language_from_url(item.get("canonical_url"))
    if language:
        return language
    title = clean(item.get("title")).lower()
    if "agentes de ia" in title or "gestión de proyectos" in title:
        return "испанском"
    if "projectbeheer" in title or "voordelen" in title:
        return "нидерландском"
    return None


def report_date(path: Path) -> datetime | None:
    match = REPORT_DATE_RE.fullmatch(path.name)
    if not match:
        return None
    return datetime.fromisoformat(match.group(1)).replace(tzinfo=timezone.utc)


def previous_report_urls(reports_dir: Path, current_date: datetime) -> set[str]:
    if not reports_dir.exists():
        return set()
    known: set[str] = set()
    current_day = current_date.date()
    for path in reports_dir.glob("AgPM_*_radar_*.md"):
        stamp = report_date(path)
        if not stamp or stamp.date() >= current_day:
            continue
        text = path.read_text(encoding="utf-8")
        for url in REPORT_LINK_RE.findall(text):
            canonical = canonicalize_report_url(url)
            if canonical:
                known.add(canonical)
                neutral = locale_neutral_report_url(canonical)
                if neutral:
                    known.add(neutral)
    return known


def event_key_from_text(text: str, canonical: str | None = None) -> str | None:
    text = clean(text).lower()
    if "meta" in text and "business" in text and "agent" in text and any(term in text for term in ["sales", "appointments", "customer service", "enterprise services", "whatsapp", "messenger", "instagram"]):
        return "event:meta_business_agent"
    if "microsoft" in text and any(term in text for term in ["project perception", "mai-cyber", "cyber-focused", "cybersecurity model", "vulnerability discovery", "enterprise cyber defense"]):
        return "event:microsoft_project_perception_security"
    if (
        "agent" in text
        and any(term in text for term in ["uniform governance", "one-size-fits-all", "one size fits all", "40%", "4 in 10", "rubbish bin", "enterprise failure"])
        and ("gartner" in text or "uniform governance" in text or "one-size" in text or "one size" in text)
    ):
        return "event:gartner_uniform_agent_governance"
    if "monday.com" in text and "ai agents" in text:
        return "event:monday_ai_agents"
    if "perplexity computer" in text:
        return "event:perplexity_computer"
    if "agent orchestration" in text or "multi-agent workflows" in text:
        return "event:agent_orchestration"
    if "agentic pmo" in text:
        return "event:agentic_pmo"
    if canonical:
        return "url:" + (locale_neutral_report_url(canonical) or canonical).lower().rstrip("/")
    return None


def text_fingerprint_key(text: str, prefix: str, min_tokens: int = 10, max_tokens: int = 16) -> str | None:
    normalized = re.sub(r"[^0-9a-zа-яё]+", " ", clean(text).lower())
    tokens = [token for token in normalized.split() if len(token) > 1]
    if len(tokens) < min_tokens:
        return None
    return f"{prefix}:" + "-".join(tokens[:max_tokens])


def habr_article_event_key(item: dict[str, Any]) -> str | None:
    source_ids = {str(hit.get("source_id") or "") for hit in item.get("source_hits", [])}
    host = urllib.parse.urlparse(str(item.get("url") or "")).netloc.lower()
    if host.startswith("www."):
        host = host[4:]
    if host not in {"habr.com", "t.me"} and not source_ids.intersection({"telegram_habr_ai", "habr_ai_hub"}):
        return None
    if host == "t.me" and "telegram_habr_ai" not in source_ids:
        return None
    return text_fingerprint_key(item_text(item), "event:habr_article")


def previous_report_event_keys(reports_dir: Path, current_date: datetime) -> set[str]:
    if not reports_dir.exists():
        return set()
    known: set[str] = set()
    current_day = current_date.date()
    for path in reports_dir.glob("AgPM_*_radar_*.md"):
        stamp = report_date(path)
        if not stamp or stamp.date() >= current_day:
            continue
        text = path.read_text(encoding="utf-8")
        for block in re.split(r"\n(?=###\s+)", text):
            if not block.startswith("### "):
                continue
            key = event_key_from_text(block)
            if key and key.startswith("event:"):
                known.add(key)
    return known


def filter_previously_reported(
    items: list[dict[str, Any]],
    reports_dir: Path,
    current_date: datetime,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    known = previous_report_urls(reports_dir, current_date)
    known_events = previous_report_event_keys(reports_dir, current_date)
    if not known:
        known = set()
    fresh: list[dict[str, Any]] = []
    skipped: list[dict[str, Any]] = []
    for item in items:
        canonical = canonicalize_report_url(item.get("canonical_url") or item.get("url"))
        neutral = locale_neutral_report_url(item.get("canonical_url") or item.get("url"))
        semantic_key = event_key(item)
        if (canonical and canonical in known) or (neutral and neutral in known) or (semantic_key.startswith("event:") and semantic_key in known_events):
            skipped.append(item)
        else:
            fresh.append(item)
    return fresh, skipped


def has_web_research_hit(item: dict[str, Any]) -> bool:
    for hit in item.get("source_hits", []):
        source_id = str(hit.get("source_id") or "")
        provider = str(hit.get("provider") or "")
        if source_id.startswith("web_") or provider in {"perplexity", "brave", "openclaw_cli"}:
            return True
    return False


def hard_missing_url(url: str | None) -> tuple[bool, int | None]:
    if not url:
        return False, None
    try:
        response = requests.get(
            url,
            headers={"User-Agent": "AgPM Radar link validator/1.0 (+https://radar.aipractice.space)"},
            allow_redirects=True,
            timeout=15,
            stream=True,
        )
        status = response.status_code
        response.close()
    except Exception:
        return False, None
    return status in {404, 410}, status


TITLE_TOKEN_STOPWORDS = {
    "the",
    "and",
    "for",
    "with",
    "from",
    "into",
    "about",
    "news",
    "для",
    "или",
    "как",
    "что",
    "это",
    "уже",
    "еще",
    "ещё",
    "при",
    "про",
    "без",
    "новости",
    "материал",
    "страница",
}


def title_tokens(value: str | None) -> set[str]:
    normalized = re.sub(r"[^0-9a-zа-яё]+", " ", clean(html.unescape(value or "")).lower())
    return {token for token in normalized.split() if len(token) > 1 and token not in TITLE_TOKEN_STOPWORDS}


def extract_title_candidates(page_text: str) -> list[str]:
    candidates: list[str] = []
    patterns = [
        r"<title[^>]*>(.*?)</title>",
        r"<meta[^>]+property=[\"']og:title[\"'][^>]+content=[\"']([^\"']+)[\"']",
        r"<meta[^>]+content=[\"']([^\"']+)[\"'][^>]+property=[\"']og:title[\"']",
        r"<h1[^>]*>(.*?)</h1>",
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, page_text[:300_000], flags=re.IGNORECASE | re.DOTALL):
            text = re.sub(r"<[^>]+>", " ", match.group(1))
            text = clean(html.unescape(text))
            if text:
                candidates.append(text)
    return candidates


def page_title_mismatch(item_title: str | None, page_text: str) -> tuple[bool, str | None]:
    item_tokens = title_tokens(item_title)
    if len(item_tokens) < 4:
        return False, None
    candidates = extract_title_candidates(page_text)
    if not candidates:
        return False, None
    page_tokens = set().union(*(title_tokens(candidate) for candidate in candidates))
    if not page_tokens:
        return False, None
    overlap = item_tokens & page_tokens
    required = max(2, min(4, round(len(item_tokens) * 0.45)))
    if len(overlap) >= required:
        return False, None
    return True, candidates[0]


def invalid_web_url(item: dict[str, Any]) -> tuple[bool, str | int | None]:
    url = item.get("url")
    if not url:
        return False, None
    try:
        response = requests.get(
            url,
            headers={"User-Agent": "AgPM Radar link validator/1.0 (+https://radar.aipractice.space)"},
            allow_redirects=True,
            timeout=15,
        )
        status = response.status_code
        if status in {404, 410}:
            response.close()
            return True, status
        content_type = response.headers.get("content-type", "")
        if "text/html" in content_type.lower():
            mismatch, page_title = page_title_mismatch(item.get("title"), response.text)
            if mismatch:
                response.close()
                return True, f"title_mismatch:{page_title}"
        response.close()
    except Exception:
        return False, None
    return False, None


def filter_hard_missing_web_links(items: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    fresh: list[dict[str, Any]] = []
    skipped: list[dict[str, Any]] = []
    for item in items:
        if not has_web_research_hit(item):
            fresh.append(item)
            continue
        invalid, status = invalid_web_url(item)
        if invalid:
            item["_link_status"] = status
            skipped.append(item)
        else:
            fresh.append(item)
    return fresh, skipped


def fulltext_cache_path(wiki: Path, url: str | None) -> Path:
    canonical = canonicalize_report_url(url) or clean(url)
    digest = hashlib.sha256(canonical.encode("utf-8")).hexdigest()[:24]
    return wiki / "data" / FULLTEXT_CACHE_DIRNAME / f"{digest}.json"


def extract_main_text(markup: str) -> str:
    soup = BeautifulSoup(markup, "html.parser")
    for node in soup(["script", "style", "noscript", "svg", "nav", "header", "footer", "form"]):
        node.decompose()
    selectors = [
        "article",
        "main",
        ".post__text",
        ".post-content",
        ".entry-content",
        ".article-content",
        ".news-detail",
        "[itemprop='articleBody']",
    ]
    chunks: list[str] = []
    for selector in selectors:
        for node in soup.select(selector):
            text = clean(node.get_text(" "))
            if len(text) >= 300 and text not in chunks:
                chunks.append(text)
        if chunks:
            break
    if not chunks and soup.body:
        chunks.append(clean(soup.body.get_text(" ")))
    return clean(" ".join(chunks))


def relevant_sentences(text: str, limit: int = 5) -> str:
    sentences = [clean(part) for part in re.split(r"(?<=[.!?])\s+", text) if clean(part)]
    scored: list[tuple[int, int, str]] = []
    terms = AGENT_TERMS + MANAGEMENT_TERMS + NEAR_TERMS + MIDDLE_TERMS + FAR_TERMS + ACTION_TERMS
    for index, sentence in enumerate(sentences):
        lowered = sentence.lower()
        score = count_terms(lowered, terms)
        if score:
            scored.append((score, -index, sentence))
    chosen = [row[2] for row in sorted(scored, reverse=True)[:limit]]
    if not chosen:
        chosen = sentences[: min(limit, len(sentences))]
    return clean(" ".join(chosen))[:4000]


def fetch_fulltext(item: dict[str, Any], wiki: Path) -> dict[str, Any] | None:
    url = item.get("url")
    if not url or not urllib.parse.urlsplit(url).scheme.startswith("http"):
        return None
    cache = fulltext_cache_path(wiki, url)
    if cache.exists():
        try:
            return json.loads(cache.read_text(encoding="utf-8"))
        except Exception:
            pass
    payload: dict[str, Any] = {
        "url": url,
        "canonical_url": canonicalize_report_url(url),
        "status": "unresolved",
        "text": "",
        "excerpt": "",
    }
    try:
        response = requests.get(
            url,
            headers={"User-Agent": "AgPM Radar fulltext extractor/1.0 (+https://radar.aipractice.space)"},
            timeout=20,
        )
        payload["http_status"] = response.status_code
        payload["content_type"] = response.headers.get("content-type", "")
        response.raise_for_status()
        if "text/html" not in payload["content_type"].lower():
            payload["status"] = "unsupported_content_type"
        else:
            text = extract_main_text(response.text)
            payload["text"] = text[:20000]
            payload["excerpt"] = relevant_sentences(text)
            payload["status"] = "resolved" if len(text) >= 300 else "weak_text"
    except Exception as exc:
        payload["error"] = str(exc)
    cache.parent.mkdir(parents=True, exist_ok=True)
    cache.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return payload


def should_fulltext_review(item: dict[str, Any], review: dict[str, Any]) -> bool:
    if source_hits(item, "ai_agents_directory_daily"):
        return False
    if review.get("perimeter") == "near":
        return True
    return int(review.get("score") or 0) >= 13


def enrich_with_fulltext_second_pass(items: list[dict[str, Any]], wiki: Path) -> tuple[list[dict[str, Any]], dict[str, int]]:
    stats = {"checked": 0, "resolved": 0, "changed": 0}
    enriched: list[dict[str, Any]] = []
    for original in items:
        item = dict(original)
        first_review = relevance_review(item)
        if not should_fulltext_review(item, first_review):
            enriched.append(item)
            continue
        stats["checked"] += 1
        payload = fetch_fulltext(item, wiki)
        if not payload or payload.get("status") not in {"resolved", "weak_text"} or not payload.get("excerpt"):
            enriched.append(item)
            continue
        stats["resolved"] += 1
        before = (first_review.get("verdict"), first_review.get("perimeter"))
        full_excerpt = clean(str(payload.get("excerpt") or ""))
        source_text = clean(str(payload.get("text") or ""))
        summary = clean(item.get("summary"))
        if full_excerpt and full_excerpt not in summary:
            item["summary"] = clean(f"{summary} По полному тексту первоисточника: {full_excerpt}") if summary else full_excerpt
        item["raw_excerpt"] = source_text or full_excerpt
        item["_fulltext_status"] = payload.get("status")
        item["_fulltext_reviewed"] = True
        second_review = relevance_review(item)
        after = (second_review.get("verdict"), second_review.get("perimeter"))
        if after != before:
            stats["changed"] += 1
            item["_fulltext_review_change"] = {
                "from": {"verdict": before[0], "perimeter": before[1]},
                "to": {"verdict": after[0], "perimeter": after[1]},
            }
        enriched.append(item)
    return enriched, stats


def has_source_hit(item: dict[str, Any], source_id: str) -> bool:
    return any(hit.get("source_id") == source_id for hit in item.get("source_hits", []))


def has_manual_source_hit(item: dict[str, Any]) -> bool:
    return any(str(hit.get("source_id") or "").startswith("manual_") for hit in item.get("source_hits", []))


def in_period(item: dict[str, Any], since: datetime, until: datetime) -> bool:
    published = parse_dt(item.get("published_at"))
    seen = parse_dt(item.get("first_seen_at")) or parse_dt(item.get("last_seen_at"))
    if has_source_hit(item, "ai_agents_directory_daily"):
        return bool((seen and since <= seen <= until) or (published and since <= published <= until))
    if has_manual_source_hit(item):
        last_seen = parse_dt(item.get("last_seen_at"))
        return bool(
            (last_seen and since <= last_seen <= until)
            or (seen and since <= seen <= until)
            or (published and since <= published <= until)
        )
    marker = published or seen
    if not marker:
        return True
    return since <= marker <= until


def clean(value: str | None) -> str:
    if not value:
        return ""
    value = re.sub(r"<<<EXTERNAL_UNTRUSTED_CONTENT[^>]*>>>", " ", value)
    value = re.sub(r"<<<END_EXTERNAL_UNTRUSTED_CONTENT[^>]*>>>", " ", value)
    value = re.sub(r"\bSource:\s*Web Search\s*---", " ", value)
    value = re.sub(r"\s*Контекст daily digest:.*$", " ", value, flags=re.IGNORECASE)
    value = re.sub(r"<<\s*>>", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def cyrillic_count(value: str) -> int:
    return len(re.findall(r"[А-Яа-яЁё]", value))


def latin_count(value: str) -> int:
    return len(re.findall(r"[A-Za-z]", value))


def needs_russian_summary(value: str) -> bool:
    if not value:
        return True
    return latin_count(value) > cyrillic_count(value) * 2


def russian_summary(item: dict[str, Any]) -> str:
    summary = clean(item.get("summary"))
    language = detected_language(item)
    item_url = (item.get("url") or "").lower()
    if "atlassian.com" in item_url and "project-management/ai-agents" in item_url and language:
        return (
            f"Материал Atlassian на {language} языке описывает AI-агентов в управлении проектами: автоматизацию задач, "
            "приоритизацию работы, поддержку решений и координацию проектной деятельности. По смыслу это локализованная версия "
            "материала о применении AI-агентов в project management; такие языковые версии нужно проверять на дубли с базовой страницей."
        )
    if summary.startswith("Материал на английском языке.") and language and language != "английском":
        return summary.replace("английском", language, 1)
    if summary and not needs_russian_summary(summary):
        return summary

    text = item_text(item)
    title = clean(item.get("title")).lower()
    url = (item.get("url") or "").lower()

    if "clawbud" in text:
        return (
            "Материал рассматривает агентную операционную среду для OpenClaw: когда агент получает браузер, каналы, файлы, CRM-данные и интеграции, "
            "ключевым становится вопрос границ, разрешений, контроля действий и управляемости агентной работы."
        )
    if "perplexity computer" in text or "openclaw" in title:
        return (
            "Материал описывает агентную рабочую среду Perplexity Computer / OpenClaw: агент может вести длинные рабочие цепочки, "
            "искать информацию, планировать, писать, запускать действия в браузере и работать с задачами. Для AgPM это сигнал о переходе "
            "от ассистента к исполняющей агентной среде."
        )
    if "monday.com" in text:
        return (
            "monday.com объявила инфраструктуру для AI-агентов внутри рабочей платформы: агенты смогут проходить аутентификацию, "
            "обновлять workflows, организовывать проекты, запускать автоматизации, формировать отчёты и координировать работу команд."
        )
    if "respond.io" in text or "respond.canny.io" in url:
        return (
            "Respond.io выпустила AI Agents для клиентских коммуникаций: агенты обрабатывают обращения от первого контакта до передачи оператору, "
            "отвечают клиентам, выполняют действия и закрывают задачи в сценариях продаж и поддержки."
        )
    if "meta" in text and "business agent" in text:
        return (
            "Meta представила Business Agent для WhatsApp, Messenger и Instagram: агент автоматизирует продажи, поддержку, запись на встречи "
            "и часть клиентских операций. Это пример бизнес-агента, который берёт на себя не только диалог, но и операционное действие."
        )
    if "agent orchestration" in text or "multi-agent workflows" in text:
        return (
            "Материал объясняет оркестрацию AI-агентов в компании: чем сложнее многоагентные workflow, тем важнее контроль согласованности действий, "
            "политики доступа, трассировка, наблюдаемость и предотвращение конфликтов между агентами."
        )
    if "letsdatascience" in url or "fight for orchestration" in title:
        return (
            "Материал фиксирует сдвиг корпоративного интереса от отдельных моделей к control plane для агентов: оркестрации, политике доступа, "
            "логированию действий, observability и безопасному подключению инструментов и данных."
        )
    if "agentic pmo" in text:
        return (
            "Материал описывает переход PMO к agentic-модели: автономные workflow помогают планировать, назначать задачи, балансировать загрузку, "
            "обнаруживать риски и готовить отчётность с меньшим объёмом ручной координации."
        )
    if "kiplot" in url:
        return (
            "Материал рассматривает AI-сценарии для PMO и портфельного управления: отчётность, портфельная видимость, контроль выполнения и снижение "
            "административной нагрузки. Для AgPM важна связка AI-агента с ответственностью и человеческим контролем."
        )
    if "hpe.com" in url:
        return (
            "Статья описывает применение agentic AI в проектном управлении: непрерывный мониторинг, динамическое перепланирование, прогнозирование рисков, "
            "ранние предупреждения и рекомендации по перераспределению ресурсов."
        )
    if "gsdcouncil" in url:
        return (
            "Обзор показывает применение agentic AI в проектной работе: автоматизированный трекинг, отчётность, приоритизация задач, управление ресурсами, "
            "коммуникации и проактивное обнаружение рисков."
        )
    if "opkey" in url:
        return (
            "Страница продукта описывает PMO AI Agent для ERP/QLM-сценариев: отслеживание работ, трассируемость, взаимодействие со стейкхолдерами "
            "и снятие части рутинной нагрузки с PMO."
        )
    if "dust.tt" in url:
        return (
            "Материал описывает AI-агентов для проектного управления: обработку проектных данных, кросс-инструментальную отчётность, сводки встреч, "
            "декомпозицию задач, обнаружение рисков, планирование ресурсов и контроль дедлайнов."
        )
    if "techplustrends" in url:
        return (
            "Материал описывает переход от ручной координации к goal-driven AI-агентам, которые распределяют задачи, прогнозируют риски "
            "и помогают управлять выполнением проектных работ."
        )
    if "wrike" in url:
        return (
            "Материал описывает AI-агентов в платформе управления работой: автономные участники команды могут выполнять многошаговые workflow без постоянного "
            "ручного запроса, а agent builder позволяет настраивать такие сценарии без кода."
        )

    if summary:
        return (
            "Материал показывает связь AI-агентов с управленческим или бизнес-контуром; детали следует проверять по первоисточнику и использовать через фильтр AgPM."
        )
    return "Краткое содержание не извлечено автоматически; материал включён по признаку связи с агентным управлением и требует просмотра первоисточника."


def item_text(item: dict[str, Any]) -> str:
    return clean(" ".join([item.get("title") or "", item.get("summary") or "", item.get("raw_excerpt") or ""])).lower()


def count_terms(text: str, terms: list[str]) -> int:
    return sum(1 for term in terms if term in text)


def url_host(item: dict[str, Any]) -> str:
    host = urllib.parse.urlsplit(item.get("url") or "").netloc.lower()
    return host[4:] if host.startswith("www.") else host


def url_path(item: dict[str, Any]) -> str:
    return urllib.parse.urlsplit(item.get("url") or "").path.lower()


def is_source_aggregator_page(item: dict[str, Any], title: str) -> bool:
    host = url_host(item)
    path = url_path(item).rstrip("/")
    return host == "aiagentsdirectory.com" and (
        path == "/news"
        or path.startswith("/news/ai-agents-news-brief")
        or "daily briefs" in title
        or "7-day summary" in title
        or "news brief" in title
    )


def is_promotional_landing_page(item: dict[str, Any], text: str) -> bool:
    path = url_path(item)
    if any(hint in path for hint in PROMOTIONAL_URL_HINTS):
        return True
    return count_terms(text, PROMOTIONAL_TEXT_TERMS) >= 2


def detect_perimeter(item: dict[str, Any]) -> str:
    text = item_text(item)
    declared = item.get("perimeter")
    if count_terms(text, NEAR_TERMS):
        return "near"
    if declared == "far":
        return "far"
    if "agent orchestration" in text or "multi-agent workflows" in text:
        return "middle"
    if count_terms(text, FAR_TERMS) or any(term in text for term in STRONG_BUSINESS_PRODUCTS):
        return "far"
    if count_terms(text, MIDDLE_TERMS):
        return "middle"
    return declared if declared in {"far", "middle", "near"} else "middle"


def relevance_review(item: dict[str, Any]) -> dict[str, Any]:
    text = item_text(item)
    title = clean(item.get("title")).lower()
    agent_score = count_terms(text, AGENT_TERMS)
    management_score = count_terms(text, MANAGEMENT_TERMS)
    near_score = count_terms(text, NEAR_TERMS)
    middle_score = count_terms(text, MIDDLE_TERMS)
    far_score = count_terms(text, FAR_TERMS)
    exclude_score = count_terms(text, EXCLUDE_TERMS)
    action_score = count_terms(text, ACTION_TERMS)
    weak_signal_score = count_terms(text, WEAK_SIGNAL_TERMS)
    strong_business_product = any(term in text for term in STRONG_BUSINESS_PRODUCTS)
    strong_management_context = any(term in text for term in STRONG_MANAGEMENT_CONTEXT)
    financial_agent_context = any(term in text for term in FINANCIAL_AGENT_TERMS)
    has_business_management_core = bool(strong_management_context or near_score or middle_score or strong_business_product)
    if "agent" in title and (" ai " in f" {title} " or "agentic" in title):
        agent_score += 1

    if "агентств" in text and not any(term in text for term in ["ии-агент", "агентн", "мультиагент"]):
        return {
            "verdict": "exclude",
            "reason": "Ложное совпадение по слову «агентство»: материал не про ИИ-агентов и не про агентное управление.",
            "perimeter": "exclude",
            "score": 0,
        }

    if not agent_score:
        return {
            "verdict": "exclude",
            "reason": "Нет признака агентности: материал может быть про ИИ вообще, но не про агентное управление.",
            "perimeter": "exclude",
            "score": 0,
        }

    if "product features" in title and not any(term in title for term in ["project management", "pmo", "portfolio"]):
        return {
            "verdict": "exclude",
            "reason": "Широкий обзор продуктовых AI-фич: проектное управление упомянуто как один из примеров, а не как агентный управленческий контур.",
            "perimeter": "exclude",
            "score": 0,
        }

    if "daily ai agent news" in title or "last 7 days" in title or is_source_aggregator_page(item, title):
        return {
            "verdict": "exclude",
            "reason": "Агрегатор новостей используется как поисковый сигнал, но не как самостоятельный материал финального радара.",
            "perimeter": "exclude",
            "score": 0,
        }

    if is_promotional_landing_page(item, text) and not near_score:
        return {
            "verdict": "exclude",
            "reason": "Вендорская продуктовая или demo-страница: есть маркетинговая упаковка решения, но нет самостоятельного аналитического или новостного материала для AgPM.",
            "perimeter": "exclude",
            "score": 0,
        }

    if ("workshop" in text or "consulting" in text) and agent_score <= 1 and "ai agent" not in title:
        return {
            "verdict": "exclude",
            "reason": "Консалтинговый или обучающий материал по AI/PMO без достаточного признака агентного управления.",
            "perimeter": "exclude",
            "score": 0,
        }

    if weak_signal_score and not has_business_management_core:
        return {
            "verdict": "exclude",
            "reason": "Материал объясняет мультиагентность или ИИ как технологию, но не показывает агентное управление в бизнесе.",
            "perimeter": "exclude",
            "score": 0,
        }

    if not has_business_management_core:
        return {
            "verdict": "exclude",
            "reason": "Есть упоминание агентов, но нет бизнес-управленческого контура: риск agent wash.",
            "perimeter": "exclude",
            "score": agent_score,
        }

    enterprise_workflow_context = bool(
        middle_score
        or strong_management_context
        or strong_business_product
        or any(term in text for term in ["workflow", "operations", "governance", "operating model", "enterprise"])
    )

    if financial_agent_context and not enterprise_workflow_context and not any(term in text for term in ["project management", "pmo", "project portfolio", "управление проект", "проектный офис"]):
        return {
            "verdict": "exclude",
            "reason": "Материал про финансовые/потребительские агентные действия, а не про агентное управление бизнесом или проектами.",
            "perimeter": "exclude",
            "score": agent_score + management_score,
        }

    if exclude_score and not (near_score >= 2 or strong_management_context or strong_business_product):
        return {
            "verdict": "exclude",
            "reason": "Материал в основном про программирование, контент, обучение или потребительское применение, а не про управление.",
            "perimeter": "exclude",
            "score": agent_score + management_score - exclude_score,
        }

    perimeter = detect_perimeter(item)
    score = agent_score * 2 + management_score + near_score * 2 + middle_score + far_score + action_score
    if exclude_score:
        score -= min(exclude_score, 3)

    if near_score and (action_score or strong_management_context or "agentic pmo" in text or "ai agents for project management" in text):
        verdict = "core"
        reason = "Прямое попадание: агентность связана с проектным управлением, PMO, портфелями или управленческой ответственностью."
    elif near_score:
        verdict = "adjacent"
        reason = "Смежно: есть проектная лексика, но слабее виден контур автономного агентного управления."
    elif middle_score or strong_business_product:
        verdict = "core" if score >= 7 and (action_score or strong_business_product or strong_management_context) else "adjacent"
        reason = "Релевантно: материал показывает агентность в бизнес-процессах, операционной модели или корпоративном управлении."
    else:
        verdict = "adjacent"
        reason = "Смежно: материал важен как сигнал по бизнес-агентам, но требует осторожной интерпретации для AgPM."

    return {"verdict": verdict, "reason": reason, "perimeter": perimeter, "score": score}


def agpm_comment(item: dict[str, Any], review: dict[str, Any]) -> str:
    text = item_text(item)
    perimeter = review["perimeter"]
    if perimeter == "near":
        return (
            "Это нужно читать как сигнал для близкого периметра AgPM: агентный слой начинает входить в задачи PMO, "
            "проектной координации, портфельной видимости или управленческого контроля. Для методологии важны не сами функции ИИ, "
            "а границы автономии, ответственность РП/PMO, журналирование решений и проверка качества агентных рекомендаций."
        )
    if perimeter == "far":
        return (
            "Это сигнал дальнего периметра: рынок движется от чат-ассистентов к агентным рабочим средам и бизнес-агентам. "
            "Для AgPM важно отслеживать, какие функции становятся исполняемыми агентами, а какие остаются интерфейсной упаковкой без управленческой автономии."
        )
    if "permission" in text or "policy" in text or "compliance" in text or "risk" in text or "governance" in text:
        return (
            "Главный вывод для AgPM — усиливается значение governance-слоя: права агента, политика доступа, контроль действий, "
            "человек в контуре и трассировка решений становятся не технической деталью, а частью управленческой архитектуры."
        )
    if "workflow" in text or "business process" in text or "operations" in text or "customer support" in text or "sales" in text:
        return (
            "Материал относится к среднему периметру: агенты встраиваются в бизнес-процессы и берут на себя не отдельную подсказку, "
            "а участок операционного исполнения. Для AgPM это важно как аналогия к проектным процессам: статус, поручения, риски, "
            "эскалации и межфункциональная координация могут становиться агентными контурами."
        )
    return (
        "Материал оставлен как смежный сигнал, но его нельзя напрямую переносить в AgPM без проверки: нужна связь с управленческим действием, "
        "ответственностью, процессом или проектным контуром."
    )


def filter_for_report(items: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    included: list[dict[str, Any]] = []
    excluded: list[dict[str, Any]] = []
    for item in items:
        review = relevance_review(item)
        item = {**item, "_radar_review": review}
        if review["verdict"] == "exclude" or review["verdict"] == "adjacent":
            if review["verdict"] == "adjacent":
                review = {
                    **review,
                    "reason": "Сигнал оставлен в базе, но не включён в финальный документ: связь с агентным управлением недостаточно сильная.",
                }
                item["_radar_review"] = review
            excluded.append(item)
        else:
            included.append(item)
    included = merge_duplicate_events(included)
    included.sort(
        key=lambda row: (
            {"core": 2, "adjacent": 1}.get(row["_radar_review"]["verdict"], 0),
            int(row["_radar_review"]["score"]),
            int(row.get("source_count", 0)),
            parse_dt(row.get("published_at")) or parse_dt(row.get("first_seen_at")) or datetime.min.replace(tzinfo=timezone.utc),
        ),
        reverse=True,
    )
    return included, excluded


def deferred_queue_path(wiki: Path) -> Path:
    return wiki / "data" / DAILY_DEFERRED_QUEUE


def load_deferred_queue(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    items: list[dict[str, Any]] = []
    with path.open("r", encoding="utf-8") as stream:
        for line in stream:
            if not line.strip():
                continue
            try:
                row = json.loads(line)
            except Exception:
                continue
            if isinstance(row, dict):
                items.append(row)
    return items


def write_deferred_queue(path: Path, items: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not items:
        path.write_text("", encoding="utf-8")
        return
    lines = [json.dumps(item, ensure_ascii=False, sort_keys=True) for item in items]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def merge_deferred_with_period(deferred: list[dict[str, Any]], period_items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: dict[str, dict[str, Any]] = {}
    for item in period_items:
        merged[event_key(item)] = item
    for item in deferred:
        merged[event_key(item)] = item
    return list(merged.values())


def clean_runtime_fields(item: dict[str, Any]) -> dict[str, Any]:
    return {key: value for key, value in item.items() if not key.startswith("_") or key == "_radar_deferred"}


def mark_deferred(item: dict[str, Any], issue_date: datetime) -> dict[str, Any]:
    clean_item = clean_runtime_fields(item)
    previous = clean_item.get("_radar_deferred") if isinstance(clean_item.get("_radar_deferred"), dict) else {}
    deferred = {
        "first_deferred_for": previous.get("first_deferred_for") or issue_date.date().isoformat(),
        "last_deferred_for": issue_date.date().isoformat(),
        "defer_count": int(previous.get("defer_count") or 0) + 1,
    }
    clean_item["_radar_deferred"] = deferred
    return clean_item


def select_daily_batch(included: list[dict[str, Any]], limit: int, issue_date: datetime) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    if limit <= 0 or len(included) <= limit:
        return included, []
    order_index = {id(item): index for index, item in enumerate(included)}

    def priority(item: dict[str, Any]) -> tuple[int, str, int]:
        deferred = item.get("_radar_deferred") if isinstance(item.get("_radar_deferred"), dict) else None
        if deferred:
            return (0, str(deferred.get("first_deferred_for") or ""), order_index.get(id(item), 0))
        return (1, "", order_index.get(id(item), 0))

    ordered = sorted(included, key=priority)
    selected = ordered[:limit]
    deferred = [mark_deferred(item, issue_date) for item in ordered[limit:]]
    return selected, deferred


def ru_materials_count(count: int) -> str:
    tail = count % 100
    last = count % 10
    if 11 <= tail <= 14:
        word = "материалов"
    elif last == 1:
        word = "материал"
    elif 2 <= last <= 4:
        word = "материала"
    else:
        word = "материалов"
    return f"{count} {word}"


def source_titles(item: dict[str, Any]) -> list[str]:
    titles = []
    for hit in item.get("source_hits", []):
        title = hit.get("source_title") or hit.get("source_id")
        if title and title not in titles:
            titles.append(title)
    return titles


def source_hits(item: dict[str, Any], source_id: str) -> list[dict[str, Any]]:
    return [hit for hit in item.get("source_hits", []) if hit.get("source_id") == source_id]


def aiagents_digest_items(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    digest_items = [item for item in items if source_hits(item, "ai_agents_directory_daily")]
    return sorted(
        digest_items,
        key=lambda row: (
            parse_dt(row.get("published_at")) or parse_dt(row.get("first_seen_at")) or datetime.min.replace(tzinfo=timezone.utc),
            clean(row.get("title")),
        ),
        reverse=True,
    )


def aiagents_digest_source_url(item: dict[str, Any]) -> str:
    for hit in source_hits(item, "ai_agents_directory_daily"):
        if hit.get("source_url"):
            return hit["source_url"]
    return "https://aiagentsdirectory.com/news"


def aiagents_digest_title(item: dict[str, Any]) -> str:
    for hit in source_hits(item, "ai_agents_directory_daily"):
        if hit.get("query"):
            return hit["query"]
    return "последний daily digest"


def aiagents_summary_parts(item: dict[str, Any]) -> tuple[str, str]:
    summary = clean(item.get("summary"))
    summary = re.sub(r"\s*Контекст daily digest:.*$", "", summary).strip()
    if "Why it matters:" in summary:
        facts, why = summary.split("Why it matters:", 1)
        return clean(facts), clean(why)
    return summary, ""


def aiagents_focus_text(item: dict[str, Any]) -> str:
    facts, why = aiagents_summary_parts(item)
    return clean(" ".join([item.get("title") or "", item.get("url") or "", item.get("raw_excerpt") or "", facts, why])).lower()


def aiagents_material_summary(item: dict[str, Any]) -> list[str]:
    text = aiagents_focus_text(item)
    title = clean(item.get("title"))
    facts, why = aiagents_summary_parts(item)

    if "tines" in text and "3b" in text:
        return [
            "Tines представила платформу Tines 3B для управления корпоративными AI-workflows, приложениями и агентами. Смысл материала — не в запуске ещё одного ИИ-инструмента, а в попытке закрыть управленческую проблему: как контролировать сложные агентные сценарии, которые начинают действовать внутри бизнес-процессов.",
            "В статье акцент сделан на governance: контроль развёртывания, управляемость AI-процессов, снижение риска несанкционированных действий и более прозрачное сопровождение корпоративных агентных контуров. Это сигнал, что рынок уже обсуждает не только «что агент умеет», но и кто задаёт ему границы, как отслеживаются действия и как организация удерживает контроль.",
        ]
    if "snowflake" in text and "cortex ai gateway" in text:
        return [
            "Материал описывает Snowflake Cortex AI Gateway — слой контроля для корпоративных AI-агентов, доступа к данным и расходов на ИИ. Важная часть новости — интеграции с партнёрами вроде 1Password и SailPoint, то есть связка агентного ИИ с управлением идентичностями, доступами и корпоративной безопасностью.",
            "Содержание материала показывает сдвиг от экспериментов с агентами к инфраструктурному управлению: предприятиям нужно понимать, какие данные агент использует, какие действия запускает и сколько стоит его работа. Это уже не витрина возможностей, а попытка встроить агентов в контур IT-governance и FinOps.",
        ]
    if "opsera" in text and "forge" in text:
        return [
            "Opsera вывела Forge в Cursor Marketplace, чтобы дать AI-assisted development средам корпоративный контекст и governance. Материал в основном относится к разработке, но его управленческий смысл — в том, что даже coding agents требуют правил, контекста предприятия и контроля исполнения.",
            "Новость показывает, что агентные инструменты разработки быстро переходят от одиночной продуктивности разработчика к управляемой фабрике ПО. В центре оказывается не сам Cursor и не генерация кода, а вопрос: как организация задаёт политикам, агентам и командам общий контекст, трассировку и правила использования.",
        ]
    if "diagrid" in text and "catalyst" in text:
        return [
            "Diagrid Catalyst 2.0 добавляет durable execution, восстановление после сбоев и подписанные истории исполнения для агентных фреймворков вроде LangGraph. В практическом смысле материал про надёжность длинных агентных процессов: агент должен не просто запуститься, а корректно продолжить работу после ошибки, сбоя среды или перезапуска.",
            "Для корпоративного применения это важнее, чем выглядит на уровне технологии. Если агент выполняет многошаговый процесс, организация должна видеть, что именно было сделано, где произошёл сбой, можно ли доказать историю исполнения и как безопасно возобновить процесс без потери управленческого контекста.",
        ]
    if "orca ade" in text or "parallel ai coding" in text:
        return [
            "Материал про Orca ADE описывает среду, где несколько coding agents могут работать параллельно без конфликтов. Основная тема — координация параллельного агентного исполнения: как не допустить, чтобы несколько агентов одновременно меняли один и тот же контур и создавали взаимные помехи.",
            "Хотя кейс относится к разработке, он полезен как модель проблемы для любых агентных процессов. Чем больше автономных исполнителей работает одновременно, тем важнее механизмы блокировок, согласования изменений, владения объектами, журналирования и разрешения конфликтов между агентными действиями.",
        ]
    if "browseract" in text:
        return [
            "BrowserAct Agent представлен как агент для создания самопроверяющихся web scrapers по текстовому запросу. Пользователь описывает нужные данные, а агент строит и тестирует скрейпер, снижая ручную работу по настройке извлечения данных.",
            "С управленческой точки зрения это материал про агентное извлечение информации из внешней среды. Он показывает, что агенты начинают брать на себя не только анализ уже подготовленных данных, но и сбор данных, проверку работоспособности сценария и поддержку повторяемого информационного процесса.",
        ]
    if "model context protocol" in text or "mcp" in text:
        return [
            "Материал описывает крупное обновление Model Context Protocol: акцент сделан на stateless-подходе, безопасности, масштабируемости и управляемости агентных подключений. MCP здесь выступает не просто техническим протоколом, а способом стандартизировать, как агент получает доступ к инструментам и данным.",
            "Главный смысл новости — зрелость инфраструктуры вокруг агентов. Когда агент подключается к корпоративным системам, критичными становятся не только качество модели, но и контекст, права, изоляция, контроль сессий, проверяемость действий и возможность масштабировать подключение без расползания рисков.",
        ]
    if "microsoft" in text and ("security" in text or "cyber" in text):
        return [
            "Microsoft представила новые AI security initiatives, включая агентную систему для киберзащитников и специализированную cyber-focused AI-модель. Материал относится к кибербезопасности, но важен тем, что показывает применение агентности в высокорисковом контуре, где ошибки могут иметь прямые организационные последствия.",
            "В таких сценариях агентность ценна за скорость обнаружения, triage, подготовку действий и сопровождение защитника. Но материал одновременно подчёркивает ограничение: чем выше риск и критичность домена, тем строже должны быть контроль, аудит, роли, права и человек в контуре принятия решений.",
        ]
    if facts:
        first = f"Материал «{title}» описывает новый агентный продукт, обновление или инфраструктурную возможность: {facts.rstrip('.')}"
        second = (
            f"Практический смысл новости: {why.rstrip('.')}" if why else
            "Практический смысл новости — в том, что агентные решения постепенно переходят из экспериментального слоя в корпоративный контур процессов, данных и контроля."
        )
        return [first + ".", second + "."]
    return [
        f"Материал «{title}» включён в ежедневную подборку AI Agents Directory как сигнал по рынку AI-агентов.",
        "Для содержательного использования в радаре его нужно читать через управленческий фильтр: есть ли в материале исполняемое действие, бизнес-процесс, ответственность, контроль, доступы, риск или проектный контур.",
    ]


def aiagents_agpm_analysis(item: dict[str, Any], review: dict[str, Any]) -> list[str]:
    text = aiagents_focus_text(item)

    if "durable" in text or "recovery" in text or "execution histories" in text:
        return [
            "Для AgPM здесь важна тема доказательной цепочки. Агентное действие должно быть воспроизводимым и проверяемым: кто инициировал сценарий, какой контекст получил агент, какие шаги выполнил, где остановился и на каком основании процесс был продолжен или отменён.",
            "В проектном управлении это напрямую переносится на поручения, контрольные точки, эскалации и статусные обновления. Если агент сопровождает проектный процесс, нужен не только результат в виде текста или статуса, но и история исполнения, пригодная для управленческой проверки.",
        ]
    if "parallel" in text or "conflict" in text:
        return [
            "Для AgPM это полезная аналогия к будущему многоагентному проектному офису. Несколько агентов могут одновременно работать с поручениями, рисками, календарём, ресурсами и коммуникациями; без правил владения объектами и разрешения конфликтов такая автоматизация быстро начнёт создавать управленческий шум.",
            "В методике внедрения это означает необходимость проектировать не только роли отдельных агентов, но и правила совместной работы: приоритеты, блокировки, уведомления, зоны ответственности, журнал изменений и механизм остановки при конфликте действий.",
        ]
    if "scraper" in text or "data extraction" in text:
        return [
            "Для AgPM материал важен как сигнал по автоматизации информационного обеспечения. Агент может сам собирать внешние данные, проверять сценарий получения информации и поддерживать повторяемый контур мониторинга, что полезно для проектной разведки, портфельных обзоров и раннего выявления изменений среды.",
            "Но такой сценарий требует аккуратной рамки доверия к данным. В проектном управлении нельзя принимать агентно собранную информацию как готовую истину: нужны источники, дата получения, проверка качества, указание неопределённости и ответственность человека за управленческий вывод.",
        ]
    if "model context protocol" in text or "mcp" in text:
        return [
            "Для AgPM это важный инфраструктурный сигнал: агентный проектный офис будет зависеть не только от модели, но и от стандартизированного доступа к инструментам, данным и контексту. MCP-подобный слой задаёт способ подключения агента к рабочей среде, а значит становится частью управленческой архитектуры.",
            "Практический вывод — в методике внедрения нужно отдельно описывать контекстный контур агента: какие системы он видит, какие действия может запускать, какие данные получает, как ограничивается сессия и где фиксируется след. Без этого агентность превращается в непрозрачную интеграцию.",
        ]
    if "microsoft" in text and ("security" in text or "cyber" in text):
        return [
            "Для AgPM материал важен как пример агентности в высокорисковой среде. Кибербезопасность хорошо показывает общий принцип: агент может ускорять обнаружение, первичную классификацию, подготовку реакции и сопровождение эксперта, но не должен незаметно подменять ответственного человека в критическом решении.",
            "Для проектного управления это переносится на риски, эскалации и изменения плана. Чем выше цена ошибки, тем строже должны быть границы автономии, журналирование, проверка источников, режим подтверждения и право человека остановить агентный сценарий.",
        ]
    if "coding" in text or "development" in text or "cursor" in text:
        return [
            "Непосредственно для AgPM это не близкий материал, потому что основной контур связан с разработкой. Но он полезен как технологический аналог: агентные исполнители уже требуют контекста организации, правил доступа, контроля изменений и координации с человеческой командой.",
            "В проектном управлении тот же принцип применим к агентам PMO: нельзя давать агенту абстрактную автономию. Ему нужна ограниченная управленческая практика, понятный объект работы, контрольные точки, журнал и правила передачи результата человеку.",
        ]
    if "govern" in text or "gateway" in text or "policy" in text or re.search(r"\baccess\b", text) or "security" in text:
        return [
            "Для AgPM это подтверждает базовую рамку: агентное управление нельзя сводить к «умному помощнику». Нужен слой управляемости — права, политики, журнал действий, контроль доступа к данным, ограничение автономии и понятная ответственность человека за итоговое решение.",
            "В прикладной методике такие материалы стоит использовать как аргументы в пользу проектирования agent governance заранее. Если в ИСУП появляется агент, который читает проектные данные, формирует выводы или запускает действия, его нужно описывать как управленческую практику с границами, а не как свободный чат-интерфейс.",
        ]
    if review["perimeter"] == "near":
        return [
            "Для AgPM это материал близкого периметра: он помогает уточнять, какие именно проектные функции могут становиться агентными — статус, контроль исполнения, риск-сканирование, подготовка решений или портфельная видимость.",
            "Главная методическая граница остаётся прежней: агент может готовить данные, варианты и действия, но ответственность за управленческое решение, изменение плана, эскалацию и распределение ресурсов должна оставаться у человека или явно назначенной роли.",
        ]
    return [
        "Для AgPM это сигнал среднего или дальнего периметра: рынок развивает инфраструктуру, безопасность и операционную надёжность AI-агентов быстрее, чем появляются зрелые сценарии именно проектного управления.",
        "Использовать такой материал стоит как слой операционализации. Он помогает формулировать требования к агентному проектному офису: управляемые права, трассировка, доказательная цепочка, наблюдаемость, контроль стоимости и человек в контуре для значимых решений.",
    ]


def render_aiagents_digest_section(items: list[dict[str, Any]]) -> list[str]:
    digest_items = aiagents_digest_items(items)
    lines: list[str] = [
        "",
        "## Разбор AI Agents Directory",
        "",
    ]
    if not digest_items:
        lines.append("За период источник AI Agents Directory не дал новых daily headline-ссылок.")
        return lines

    first = digest_items[0]
    lines.extend(
        [
            f"Источник: {aiagents_digest_source_url(first)}",
            f"Daily digest: {aiagents_digest_title(first)}",
            f"Проверено внешних ссылок из подборки: {ru_materials_count(len(digest_items))}.",
            "",
            "Каждая ссылка из daily-подборки проходит тот же смысловой фильтр, что и остальные материалы радара. В общий радар проходят только материалы с управленческим контуром; остальные фиксируются как отсечённые или смежные сигналы.",
            "",
        ]
    )
    for item in digest_items:
        review = relevance_review(item)
        lines.extend(
            [
                f"### {clean(item.get('title')) or 'Без названия'}",
                "",
                f"Ссылка: {item.get('url', '')}",
                "",
                "Суть материала:",
                "",
                *aiagents_material_summary(item),
                "",
                "Вывод для AgPM:",
                "",
                *aiagents_agpm_analysis(item, review),
                "",
            ]
        )
    return lines


def summary_sentences(value: str, limit: int = 2) -> list[str]:
    value = clean(value)
    if not value:
        return []
    parts = re.split(r"(?<=[.!?])\s+", value)
    sentences = [part.strip() for part in parts if part.strip()]
    if len(sentences) <= limit:
        return sentences
    return sentences[:limit]


def general_material_summary(item: dict[str, Any]) -> list[str]:
    text = item_text(item)
    title = clean(item.get("title"))
    summary = clean(item.get("summary"))
    ru_summary = russian_summary(item)

    if summary and not needs_russian_summary(summary) and "детали следует проверять" not in summary:
        sentences = summary_sentences(summary, 2)
        first = " ".join(sentences) if sentences else summary
    elif "workflow" in text or "business process" in text or "orchestration" in text:
        first = (
            f"Материал «{title}» описывает переход от отдельных AI-помощников к агентным workflow: агенты получают роль в последовательности действий, "
            "координации задач, обработке данных и сопровождении бизнес-процесса."
        )
    elif "pmo" in text or "project management" in text or "portfolio" in text:
        first = (
            f"Материал «{title}» относится к близкому периметру: в нём агентность связана с проектным управлением, PMO, портфельной видимостью, "
            "контролем исполнения или снижением ручной управленческой нагрузки."
        )
    elif "governance" in text or "policy" in text or "permission" in text or "compliance" in text or "risk" in text:
        first = (
            f"Материал «{title}» рассматривает управляемость AI-агентов: политики, права, контроль действий, риски, соответствие требованиям и "
            "корпоративную эксплуатацию агентных сценариев."
        )
    else:
        first = ru_summary

    if "Материал показывает связь AI-агентов" in first or "Краткое содержание не извлечено" in first:
        first = (
            f"Материал «{title}» зафиксирован как релевантный сигнал по рынку AI-агентов. Его ценность для радара не в самой новости об ИИ, "
            "а в наличии связи с бизнес-процессом, управлением, ответственностью, контролем или проектным контуром."
        )

    if "workflow" in text or "business process" in text or "operations" in text:
        second = (
            "Практически это показывает, что агентность смещается из интерфейса подсказок в операционное исполнение: агент начинает сопровождать участок процесса, "
            "готовить действия, связывать системы и снижать ручную координацию."
        )
    elif "pmo" in text or "project management" in text or "portfolio" in text:
        second = (
            "Для проектной среды это важно как сигнал о созревании прикладных сценариев: статусная отчётность, контроль поручений, анализ рисков, "
            "подготовка повесток и портфельные обзоры становятся естественными зонами агентной поддержки."
        )
    elif "governance" in text or "policy" in text or "permission" in text or "compliance" in text or "risk" in text:
        second = (
            "Главный акцент материала — не производительность модели, а управленческая надстройка вокруг неё: кто разрешает действие, как ограничивается доступ, "
            "как проверяется результат и где остаётся ответственность человека."
        )
    else:
        second = (
            "Использовать материал в радаре стоит через фильтр AgPM: отделять реальное делегирование управленческого действия от маркетингового agent wash "
            "и проверять, есть ли в новости процесс, роль, объект управления и доказательная цепочка."
        )
    return [first, second]


def general_agpm_analysis(item: dict[str, Any], review: dict[str, Any]) -> list[str]:
    text = item_text(item)
    perimeter = review.get("perimeter")

    if perimeter == "near":
        return [
            "Для AgPM это материал близкого периметра: он помогает уточнять, какие проектные функции можно переводить в агентный режим без подмены управленческой ответственности. В центре здесь не «ИИ вообще», а конкретные практики PMO: наблюдение, статус, риск, поручения, координация и портфельная видимость.",
            "Методически такой материал стоит использовать для операционализации AgPM: описывать границы автономии агента, точки человеческого подтверждения, журнал действий и критерии качества агентной рекомендации. Чем ближе сценарий к изменению плана, сроков, ресурсов или эскалации, тем сильнее должен быть human-in-the-loop.",
        ]
    if "governance" in text or "policy" in text or "permission" in text or "compliance" in text or "risk" in text:
        return [
            "Для AgPM материал усиливает governance-линию: агентная система должна быть управляемой, наблюдаемой и ограниченной правилами. Права, политики, журналирование, контроль доступа и трассировка действий становятся не технической деталью, а частью управленческой архитектуры.",
            "В прикладной методике это означает, что внедрение агента нужно начинать не с перечня функций, а с описания управленческой практики: что агент делает, где останавливается, кто подтверждает значимое действие, какие данные используются и как доказывается корректность результата.",
        ]
    if "workflow" in text or "business process" in text or "operations" in text or "orchestration" in text:
        return [
            "Для AgPM это материал среднего периметра: бизнес-агенты показывают, как автономное или полуавтономное действие встраивается в процесс. Проектное управление может использовать эту логику для статусов, поручений, контроля сроков, подготовки решений и межфункциональной координации.",
            "Ключевой вывод — агентность нужно проектировать как процессный контур, а не как отдельный чат. Для проектного офиса это означает связку роли агента, входных данных, допустимых действий, контрольных точек, журнала и ответственности человека за управленческий результат.",
        ]
    if perimeter == "far":
        return [
            "Для AgPM это сигнал дальнего периметра: рынок развивает агентные рабочие среды, платформы и инфраструктуру быстрее, чем появляются зрелые специализированные PMO-сценарии. Такие материалы полезны как ранние индикаторы будущих требований к агентному проектному офису.",
            "Использовать их следует осторожно: они не меняют канон AgPM напрямую, но помогают уточнять требования к инструментальному слою — доступы, память, наблюдаемость, оркестрация, стоимость агентных запусков, безопасность и возможность остановки агентного действия.",
        ]
    return [
        "Для AgPM материал полезен как смежный управленческий сигнал: он показывает, где агентность начинает переходить от генерации текста к выполнению действий в цифровом контуре организации.",
        "При переносе в проектное управление важно сохранять консервативную рамку: агент помогает наблюдать, готовить и исполнять ограниченные операции, но ответственность за значимое решение, приоритет, риск и изменение плана остаётся у человека.",
    ]


def event_key(item: dict[str, Any]) -> str:
    habr_key = habr_article_event_key(item)
    if habr_key:
        return habr_key
    canonical = item.get("canonical_url") or item.get("url") or clean(item.get("title"))
    neutral = locale_neutral_report_url(canonical) or str(canonical)
    return event_key_from_text(item_text(item), canonical) or "url:" + neutral.lower().rstrip("/")


def source_preference(item: dict[str, Any]) -> int:
    host = urllib.parse.urlparse(str(item.get("url") or "")).netloc.lower()
    if host.startswith("www."):
        host = host[4:]
    if host == "gartner.com":
        return 3
    if host == "habr.com":
        return 2
    if host in {"reuters.com", "marketing4ecommerce.net"}:
        return 2
    if host in {"theregister.com", "cloudwars.com", "techradar.com", "infosecurity-magazine.com"}:
        return 1
    return 0


def merge_duplicate_events(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: dict[str, dict[str, Any]] = {}
    for item in items:
        key = event_key(item)
        if key not in merged:
            merged[key] = item
            continue
        current = merged[key]
        current_review = current["_radar_review"]
        item_review = item["_radar_review"]
        current_rank = (
            {"core": 2, "adjacent": 1}.get(current_review["verdict"], 0),
            source_preference(current),
            int(current_review["score"]),
            int(current.get("source_count", 0)),
        )
        item_rank = (
            {"core": 2, "adjacent": 1}.get(item_review["verdict"], 0),
            source_preference(item),
            int(item_review["score"]),
            int(item.get("source_count", 0)),
        )
        if item_rank > current_rank:
            item, current = current, item
            merged[key] = current
        current_hits = current.setdefault("source_hits", [])
        known = {(hit.get("source_id"), hit.get("hit_url")) for hit in current_hits}
        for hit in item.get("source_hits", []):
            marker = (hit.get("source_id"), hit.get("hit_url"))
            if marker not in known:
                current_hits.append(hit)
                known.add(marker)
        alt_links = current.setdefault("alternative_links", [])
        if item.get("url") and item.get("url") != current.get("url") and item.get("url") not in alt_links:
            alt_links.append(item["url"])
        current["source_count"] = len({hit.get("source_id") for hit in current_hits if hit.get("source_id")})
        current["hit_count"] = len(current_hits)
    return list(merged.values())


def top_items(items: list[dict[str, Any]], limit: int = 12) -> list[dict[str, Any]]:
    return sorted(
        items,
        key=lambda row: (int(row.get("source_count", 0)), parse_dt(row.get("published_at")) or parse_dt(row.get("first_seen_at")) or datetime.min.replace(tzinfo=timezone.utc)),
        reverse=True,
    )[:limit]


def agpm_implications(items: list[dict[str, Any]]) -> list[str]:
    text = " ".join(item_text(item) for item in items).lower()
    implications: list[str] = []

    if any(word in text for word in ["computer use", "browser", "workspace", "openclaw", "perplexity computer"]):
        implications.append(
            "Усилить в AgPM различение между агентом как управленческим участником и агентной рабочей средой: новые продукты всё чаще берут на себя не только генерацию текста, но и выполнение действий в цифровом контуре."
        )
    if any(word in text for word in ["workflow", "business process", "bpm", "operating model", "enterprise"]):
        implications.append(
            "Для среднего периметра стоит уточнять связку AgPM с агентными бизнес-процессами: проектное управление должно показывать, как агентные операции входят в контур сроков, ответственности, рисков и изменений."
        )
    if any(word in text for word in ["pmo", "project management", "portfolio", "ai pmo", "проект"]):
        implications.append(
            "Близкий периметр требует отдельного watchlist по AI PMO: важно отслеживать не только инструменты, но и управленческие сценарии — контроль поручений, статусную отчётность, риск-сканирование, координацию и портфельные решения."
        )
    if not implications:
        implications.append(
            "За период не видно достаточного массива близких материалов; полезно расширить поисковые запросы и отдельно проверить профессиональные PM/PMO-источники."
        )
    implications.append(
        "Для канона AgPM новые материалы следует использовать как слой операционализации и рыночной разведки, а не как автоматическое основание для изменения принципов."
    )
    return implications


def render_markdown(
    items: list[dict[str, Any]],
    since: datetime,
    until: datetime,
    included_override: list[dict[str, Any]] | None = None,
    excluded_override: list[dict[str, Any]] | None = None,
    deferred_count: int = 0,
) -> str:
    if included_override is None or excluded_override is None:
        included, excluded = filter_for_report(items)
    else:
        included, excluded = included_override, excluded_override
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for item in included:
        grouped[item["_radar_review"]["perimeter"]].append(item)

    counts = Counter(item["_radar_review"]["perimeter"] for item in included)
    verdicts = Counter(item["_radar_review"]["verdict"] for item in included)
    period = f"{since.date().isoformat()} — {until.date().isoformat()}"
    lines: list[str] = [
        "# Еженедельный радар активностей по AgPM",
        "",
        f"Период: {period}",
        f"Просмотрено материалов: {len(items)}",
        f"Включено в радар после смыслового отбора: {len(included)}",
        f"Отсеяно как нерелевантное агентному управлению: {len(excluded)}",
        "",
        "## Оглавление",
        "",
        "Быстрый переход к четырём ключевым разделам радара:",
        "",
        *[f"- [{title}](#{anchor})" for title, anchor in TOC_LINKS],
        "",
        "## Краткое резюме",
        "",
        f"- Дальний периметр: {ru_materials_count(counts.get('far', 0))}.",
        f"- Средний периметр: {ru_materials_count(counts.get('middle', 0))}.",
        f"- Близкий периметр: {ru_materials_count(counts.get('near', 0))}.",
        f"- Ядро радара: {ru_materials_count(verdicts.get('core', 0))}; смежные сигналы: {ru_materials_count(verdicts.get('adjacent', 0))}.",
        "",
        "В финальный радар не включаются материалы, где агентность является только маркетинговой меткой, а также материалы про программирование, генерацию контента, потребительские ИИ-инструменты или общие новости ИИ без управленческого контура.",
        "",
        "## Что важно для AgPM",
        "",
    ]
    if deferred_count:
        lines.insert(7, f"Перенесено в очередь следующего daily-выпуска: {ru_materials_count(deferred_count)}")
    lines.extend([f"- {item}" for item in agpm_implications(included)])
    lines.extend(render_aiagents_digest_section(included))

    for perimeter, title, desc in PERIMETERS:
        perimeter_items = grouped.get(perimeter, [])
        section_items = top_items(perimeter_items, max(10, len(perimeter_items)))
        lines.extend(["", f"## {title}", "", desc + ".", ""])
        if not section_items:
            lines.append("За период материалов, прошедших смысловой фильтр, не найдено.")
            continue
        for item in section_items:
            review = item["_radar_review"]
            found_via = ", ".join(source_titles(item))
            lines.extend(
                [
                    f"### {clean(item.get('title')) or 'Без названия'}",
                    "",
                    f"Ссылка: {item.get('url', '')}",
                    f"Найдено через: {found_via or 'источник не указан'}",
                    "",
                    "Суть материала:",
                    "",
                    *general_material_summary(item),
                    "",
                    "Вывод для AgPM:",
                    "",
                    *general_agpm_analysis(item, review),
                    "",
                ]
            )
            if item.get("alternative_links"):
                links = "; ".join(item["alternative_links"][:3])
                lines.insert(-1, f"- Дополнительные ссылки по тому же событию: {links}")

    intersections = [item for item in included if int(item.get("source_count", 0)) > 1]
    lines.extend(["", "## Пересечения между источниками", ""])
    if intersections:
        for item in top_items(intersections, 10):
            lines.append(f"- {clean(item.get('title'))}: {item.get('source_count')} источника; {item.get('url', '')}")
    else:
        lines.append("Среди материалов, прошедших смысловой фильтр, пересечений между источниками за период не выявлено.")

    excluded_reasons = Counter(item["_radar_review"]["reason"] for item in excluded)
    lines.extend(["", "## Что отсечено", ""])
    if excluded_reasons:
        for reason, count in excluded_reasons.most_common():
            lines.append(f"- {reason}: {ru_materials_count(count)}.")
    else:
        lines.append("Отсечённых материалов нет.")

    lines.extend(
        [
            "",
            "## Методическая оговорка",
            "",
            "Радар фиксирует не общий поток ИИ-новостей, а только материалы, которые помогают понять развитие агентного управления. Материал считается релевантным, если в нём есть не только ИИ или агент как технология, но и управленческий контур: делегирование действий, бизнес-процесс, ответственность, governance, контроль, риск, PMO, портфель или проектная координация. Сигналы рынка используются как слой операционализации и разведки, а не как автоматическое основание для изменения канона AgPM.",
            "",
        ]
    )
    return "\n".join(lines)


URL_RE = re.compile(r"https?://[^;\s]+")


def style_link_run(run: Any) -> None:
    run.font.color.rgb = RGBColor(5, 99, 193)
    run.font.underline = True


def add_external_hyperlink(paragraph: Any, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_internal_hyperlink(paragraph: Any, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph: Any, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_text_with_links(paragraph: Any, text: str) -> None:
    position = 0
    for match in URL_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        raw_url = match.group(0)
        url = raw_url.rstrip(".,)")
        trailing = raw_url[len(url) :]
        add_external_hyperlink(paragraph, url, url)
        if trailing:
            paragraph.add_run(trailing)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def add_markdown_to_docx(markdown: str, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    first_heading = True
    bookmark_id = 1
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if first_heading else WD_PARAGRAPH_ALIGNMENT.LEFT
            p.add_run(line[2:].strip())
            first_heading = False
        elif line.startswith("## "):
            title = line[3:].strip()
            p = doc.add_paragraph(title, style="Heading 2")
            if title in SECTION_ANCHORS:
                add_bookmark(p, SECTION_ANCHORS[title], bookmark_id)
                bookmark_id += 1
        elif line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
        elif line.startswith("- "):
            item = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            internal = re.fullmatch(r"\[([^\]]+)\]\(#([^)]+)\)", item)
            if internal:
                add_internal_hyperlink(p, internal.group(1), internal.group(2))
            else:
                add_text_with_links(p, item)
        else:
            p = doc.add_paragraph()
            add_text_with_links(p, line)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.12
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate the AgPM radar report (weekly by default; use --output-prefix daily for daily mode).")
    parser.add_argument("--wiki", type=Path, default=DEFAULT_WIKI)
    parser.add_argument("--days", type=int, default=7)
    parser.add_argument("--until", help="UTC date YYYY-MM-DD. Defaults to today.")
    parser.add_argument("--output-dir", type=Path)
    parser.add_argument("--output-prefix", default="weekly", help="Output file prefix (weekly or daily).")
    return parser.parse_args()


def parse_until(value: str | None) -> datetime:
    if not value:
        return utc_now()
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", value):
        return datetime.fromisoformat(value).replace(hour=23, minute=59, second=59, tzinfo=timezone.utc)
    return datetime.fromisoformat(value.replace("Z", "+00:00")).astimezone(timezone.utc)


def main() -> int:
    args = parse_args()
    until = parse_until(args.until)
    since = until - timedelta(days=args.days)
    materials = load_materials(args.wiki / "data" / "materials.jsonl")
    period_items = [item for item in materials if in_period(item, since, until)]
    skipped_previous: list[dict[str, Any]] = []
    skipped_dead_links: list[dict[str, Any]] = []
    deferred_written = 0
    if args.output_prefix == "daily":
        queue_path = deferred_queue_path(args.wiki)
        period_items = merge_deferred_with_period(load_deferred_queue(queue_path), period_items)
        period_items, skipped_previous = filter_previously_reported(period_items, args.wiki / "reports", until)
        period_items, skipped_dead_links = filter_hard_missing_web_links(period_items)
        period_items, fulltext_stats = enrich_with_fulltext_second_pass(period_items, args.wiki)
        included, excluded = filter_for_report(period_items)
        included, deferred_items = select_daily_batch(included, DAILY_REPORT_LIMIT, until)
        write_deferred_queue(queue_path, deferred_items)
        deferred_written = len(deferred_items)
    else:
        fulltext_stats = {"checked": 0, "resolved": 0, "changed": 0}
        included = None
        excluded = None

    markdown = render_markdown(period_items, since, until, included, excluded, deferred_written)
    output_dir = args.output_dir or args.wiki / "reports"
    output_dir.mkdir(parents=True, exist_ok=True)
    stamp = until.date().isoformat()
    prefix = args.output_prefix
    md_path = output_dir / f"AgPM_{prefix}_radar_{stamp}.md"
    docx_path = output_dir / f"AgPM_{prefix}_radar_{stamp}.docx"
    md_path.write_text(markdown, encoding="utf-8")
    add_markdown_to_docx(markdown, docx_path)
    print(json.dumps({"ok": True, "markdown": str(md_path), "docx": str(docx_path), "items": len(period_items), "included": len(included) if included is not None else None, "deferred_next_issue": deferred_written, "skipped_previous_issues": len(skipped_previous), "skipped_dead_links": len(skipped_dead_links), "fulltext_second_pass": fulltext_stats}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
